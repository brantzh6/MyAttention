"""
Knowledge Base Manager - Multi-collection RAG system

Provides flexible knowledge management with:
- Multiple knowledge bases (collections)
- Multiple file format support
- Web search integration (Aliyun IQS)
- Document metadata management
"""

import os
import sys
from typing import List, Dict, Any, Optional, BinaryIO
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
import hashlib
import uuid
import mimetypes

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from qdrant_client.http import models as qmodels
from sentence_transformers import SentenceTransformer

# Optional imports for file parsing
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    from bs4 import BeautifulSoup
    HTML_SUPPORT = True
except ImportError:
    HTML_SUPPORT = False

try:
    import markdown
    MARKDOWN_SUPPORT = True
except ImportError:
    MARKDOWN_SUPPORT = False

try:
    import requests
    from urllib.parse import urlparse
    WEB_SUPPORT = True
except ImportError:
    WEB_SUPPORT = False

from config import create_qdrant_client, get_settings


@dataclass
class KnowledgeBase:
    """Knowledge base metadata"""
    id: str
    name: str
    description: str = ""
    created_at: datetime = field(default_factory=datetime.now)
    updated_at: datetime = field(default_factory=datetime.now)
    document_count: int = 0
    chunk_count: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class Document:
    """A document with its metadata"""
    id: str
    content: str
    title: str = ""
    source: str = ""
    source_type: str = ""
    url: str = ""
    file_type: str = ""
    file_size: int = 0
    kb_id: str = "default"
    created_at: datetime = field(default_factory=datetime.now)
    metadata: Dict[str, Any] = field(default_factory=dict)

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


@dataclass
class SearchResult:
    """A search result with relevance score"""
    document: Document
    score: float
    chunk_index: int = 0


@dataclass
class FileParseResult:
    """Result of parsing a file"""
    success: bool
    content: str = ""
    title: str = ""
    error: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)


class FileParser:
    """Parse various file formats to text"""

    SUPPORTED_EXTENSIONS = {
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.markdown': 'text/markdown',
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.html': 'text/html',
        '.htm': 'text/html',
    }

    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """Check if file extension is supported"""
        ext = Path(filename).suffix.lower()
        return ext in cls.SUPPORTED_EXTENSIONS

    @classmethod
    def parse(cls, file_content: bytes, filename: str) -> FileParseResult:
        """Parse file content based on extension"""
        ext = Path(filename).suffix.lower()

        if ext == '.txt':
            return cls._parse_txt(file_content, filename)
        elif ext in ['.md', '.markdown']:
            return cls._parse_markdown(file_content, filename)
        elif ext == '.pdf':
            return cls._parse_pdf(file_content, filename)
        elif ext == '.docx':
            return cls._parse_docx(file_content, filename)
        elif ext in ['.html', '.htm']:
            return cls._parse_html(file_content, filename)
        else:
            return FileParseResult(
                success=False,
                error=f"Unsupported file type: {ext}"
            )

    @classmethod
    def _decode_text(cls, content: bytes) -> str:
        """Try to decode bytes to string with multiple encodings"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'ascii']
        for encoding in encodings:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError("Unable to decode file with any known encoding")

    @classmethod
    def _parse_txt(cls, content: bytes, filename: str) -> FileParseResult:
        """Parse plain text file"""
        try:
            text = cls._decode_text(content)
            return FileParseResult(
                success=True,
                content=text,
                title=Path(filename).stem,
                metadata={'encoding': 'auto-detected'}
            )
        except Exception as e:
            return FileParseResult(success=False, error=str(e))

    @classmethod
    def _parse_markdown(cls, content: bytes, filename: str) -> FileParseResult:
        """Parse Markdown file"""
        try:
            text = cls._decode_text(content)
            # Extract title from first heading
            title = Path(filename).stem
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('# '):
                    title = line[2:].strip()
                    break

            result = FileParseResult(
                success=True,
                content=text,
                title=title,
                metadata={'format': 'markdown'}
            )

            # Convert to HTML if markdown library available
            if MARKDOWN_SUPPORT:
                result.metadata['html'] = markdown.markdown(text)

            return result
        except Exception as e:
            return FileParseResult(success=False, error=str(e))

    @classmethod
    def _parse_pdf(cls, content: bytes, filename: str) -> FileParseResult:
        """Parse PDF file"""
        if not PDF_SUPPORT:
            return FileParseResult(
                success=False,
                error="PDF support not installed. Run: pip install PyPDF2"
            )

        try:
            import io
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_parts = []
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())

            text = '\n\n'.join(text_parts)

            # Try to get title from metadata
            title = Path(filename).stem
            if pdf_reader.metadata and pdf_reader.metadata.title:
                title = pdf_reader.metadata.title

            return FileParseResult(
                success=True,
                content=text,
                title=title,
                metadata={
                    'pages': len(pdf_reader.pages),
                    'format': 'pdf'
                }
            )
        except Exception as e:
            return FileParseResult(success=False, error=str(e))

    @classmethod
    def _parse_docx(cls, content: bytes, filename: str) -> FileParseResult:
        """Parse DOCX file"""
        if not DOCX_SUPPORT:
            return FileParseResult(
                success=False,
                error="DOCX support not installed. Run: pip install python-docx"
            )

        try:
            import io
            doc_file = io.BytesIO(content)
            doc = DocxDocument(doc_file)

            text_parts = []
            for para in doc.paragraphs:
                text_parts.append(para.text)

            text = '\n'.join(text_parts)

            # Try to get title from core properties
            title = Path(filename).stem
            if doc.core_properties.title:
                title = doc.core_properties.title

            return FileParseResult(
                success=True,
                content=text,
                title=title,
                metadata={
                    'paragraphs': len(doc.paragraphs),
                    'format': 'docx'
                }
            )
        except Exception as e:
            return FileParseResult(success=False, error=str(e))

    @classmethod
    def _parse_html(cls, content: bytes, filename: str) -> FileParseResult:
        """Parse HTML file"""
        if not HTML_SUPPORT:
            return FileParseResult(
                success=False,
                error="HTML support not installed. Run: pip install beautifulsoup4"
            )

        try:
            text = cls._decode_text(content)
            soup = BeautifulSoup(text, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get title
            title = Path(filename).stem
            if soup.title and soup.title.string:
                title = soup.title.string.strip()

            # Get text content
            text = soup.get_text(separator='\n', strip=True)

            # Clean up whitespace
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            text = '\n'.join(lines)

            return FileParseResult(
                success=True,
                content=text,
                title=title,
                metadata={'format': 'html'}
            )
        except Exception as e:
            return FileParseResult(success=False, error=str(e))

    @classmethod
    def parse_url(cls, url: str, timeout: int = 30, verify_ssl: bool = True) -> FileParseResult:
        """Fetch and parse webpage"""
        if not WEB_SUPPORT:
            return FileParseResult(
                success=False,
                error="Web support not installed. Run: pip install requests beautifulsoup4"
            )

        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
            }

            # Disable SSL verification for development (set verify=True in production)
            import os
            verify = verify_ssl if verify_ssl else not os.getenv('DISABLE_SSL_VERIFY', 'false').lower() == 'true'

            response = requests.get(url, headers=headers, timeout=timeout, verify=verify)
            response.raise_for_status()

            # Parse as HTML
            soup = BeautifulSoup(response.content, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header", "aside", "advertisement"]):
                script.decompose()

            # Get title
            title = ""
            if soup.title and soup.title.string:
                title = soup.title.string.strip()

            # Try to get main content - prioritize article/main content areas
            main_content = (
                soup.find('article') or
                soup.find('main') or
                soup.find('div', class_=['content', 'post-content', 'article-content', 'entry-content']) or
                soup.find('section', class_=['content', 'main'])
            )

            if main_content:
                text = main_content.get_text(separator='\n', strip=True)
            else:
                # Fall back to body content, excluding common non-content areas
                body = soup.find('body')
                if body:
                    # Remove common non-content elements
                    for elem in body.find_all(['sidebar', 'comments', 'related']):
                        elem.decompose()
                    text = body.get_text(separator='\n', strip=True)
                else:
                    text = soup.get_text(separator='\n', strip=True)

            # Clean up whitespace
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            text = '\n'.join(lines)

            return FileParseResult(
                success=True,
                content=text,
                title=title or urlparse(url).netloc,
                metadata={
                    'format': 'webpage',
                    'url': url,
                    'status_code': response.status_code,
                    'content_type': response.headers.get('Content-Type', ''),
                }
            )
        except requests.exceptions.SSLError as e:
            return FileParseResult(
                success=False,
                error=f"SSL certificate verification failed. You can set DISABLE_SSL_VERIFY=true to bypass this in development. Error: {str(e)}"
            )
        except requests.exceptions.Timeout:
            return FileParseResult(
                success=False,
                error=f"Request timeout after {timeout} seconds"
            )
        except requests.exceptions.ConnectionError as e:
            return FileParseResult(
                success=False,
                error=f"Connection error: {str(e)}"
            )
        except Exception as e:
            return FileParseResult(success=False, error=str(e))


class KnowledgeBaseManager:
    """
    Multi-collection RAG Engine for knowledge management

    Features:
    - Multiple knowledge bases (collections)
    - Document chunking and embedding
    - Semantic similarity search
    - Multiple file format support
    - Web search integration
    """

    EMBEDDING_MODEL = "all-MiniLM-L6-v2"
    EMBEDDING_DIM = 384
    CHUNK_SIZE = 500
    CHUNK_OVERLAP = 50

    # Metadata collection name for storing KB info
    METADATA_COLLECTION = "kb_metadata"

    def __init__(self):
        settings = get_settings()
        self.client = create_qdrant_client(settings)
        self._encoder = None
        self._kb_cache: Dict[str, KnowledgeBase] = {}
        self._ensure_metadata_collection()

    @property
    def encoder(self) -> SentenceTransformer:
        """Lazy load the embedding model"""
        if self._encoder is None:
            self._encoder = SentenceTransformer(self.EMBEDDING_MODEL)
        return self._encoder

    def _ensure_metadata_collection(self):
        """Create metadata collection if it doesn't exist"""
        collections = self.client.get_collections().collections
        exists = any(c.name == self.METADATA_COLLECTION for c in collections)

        if not exists:
            self.client.create_collection(
                collection_name=self.METADATA_COLLECTION,
                vectors_config=qmodels.VectorParams(
                    size=2,  # Minimal size for metadata storage
                    distance=qmodels.Distance.COSINE,
                ),
            )

    def _save_kb_metadata(self, kb: KnowledgeBase):
        """Save knowledge base metadata to Qdrant"""
        try:
            # Generate a deterministic UUID from the kb_id for the point ID
            import uuid
            point_uuid = str(uuid.uuid5(uuid.NAMESPACE_DNS, kb.id))

            self.client.upsert(
                collection_name=self.METADATA_COLLECTION,
                points=[qmodels.PointStruct(
                    id=point_uuid,
                    vector=[0.0, 0.0],  # Dummy vector
                    payload={
                        "kb_id": kb.id,
                        "name": kb.name,
                        "description": kb.description,
                        "created_at": kb.created_at.isoformat(),
                        "updated_at": kb.updated_at.isoformat(),
                    },
                )],
            )
        except Exception as e:
            print(f"Warning: Failed to save KB metadata: {e}")

    def _load_kb_metadata(self, kb_id: str) -> Optional[Dict[str, Any]]:
        """Load knowledge base metadata from Qdrant"""
        try:
            import uuid
            point_uuid = str(uuid.uuid5(uuid.NAMESPACE_DNS, kb_id))

            result = self.client.retrieve(
                collection_name=self.METADATA_COLLECTION,
                ids=[point_uuid],
            )
            if result:
                return result[0].payload
        except Exception:
            pass
        return None

    def _get_collection_name(self, kb_id: str) -> str:
        """Get Qdrant collection name for a knowledge base"""
        return f"kb_{kb_id}"

    def _ensure_collection(self, kb_id: str):
        """Create collection if it doesn't exist"""
        collection_name = self._get_collection_name(kb_id)
        collections = self.client.get_collections().collections
        exists = any(c.name == collection_name for c in collections)

        if not exists:
            self.client.create_collection(
                collection_name=collection_name,
                vectors_config=qmodels.VectorParams(
                    size=self.EMBEDDING_DIM,
                    distance=qmodels.Distance.COSINE,
                ),
            )

    def _chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        if len(text) <= self.CHUNK_SIZE:
            return [text]

        chunks = []
        start = 0
        while start < len(text):
            end = start + self.CHUNK_SIZE
            chunk = text[start:end]

            # Try to break at sentence boundary
            if end < len(text):
                last_period = chunk.rfind('。')
                if last_period == -1:
                    last_period = chunk.rfind('. ')
                if last_period > self.CHUNK_SIZE // 2:
                    chunk = chunk[:last_period + 1]
                    end = start + last_period + 1

            chunks.append(chunk.strip())
            start = end - self.CHUNK_OVERLAP

        return [c for c in chunks if c]

    def _generate_id(self, content: str, source: str) -> str:
        """Generate deterministic ID from content hash"""
        hash_input = f"{source}:{content[:200]}"
        return hashlib.md5(hash_input.encode()).hexdigest()

    # ==================== Knowledge Base Management ====================

    async def create_knowledge_base(
        self,
        name: str,
        description: str = "",
        metadata: Dict[str, Any] = None
    ) -> KnowledgeBase:
        """Create a new knowledge base"""
        kb_id = str(uuid.uuid4())[:8]

        kb = KnowledgeBase(
            id=kb_id,
            name=name,
            description=description,
            metadata=metadata or {}
        )

        # Create Qdrant collection
        self._ensure_collection(kb_id)

        # Save metadata to Qdrant
        self._save_kb_metadata(kb)

        # Cache it
        self._kb_cache[kb_id] = kb

        return kb

    async def list_knowledge_bases(self) -> List[KnowledgeBase]:
        """List all knowledge bases"""
        collections = self.client.get_collections().collections
        kbs = []

        for col in collections:
            # Skip metadata collection and non-kb collections
            if not col.name.startswith("kb_") or col.name == self.METADATA_COLLECTION:
                continue

            kb_id = col.name[3:]  # Remove "kb_" prefix

            # Get stats
            info = self.client.get_collection(col.name)

            # Load metadata
            metadata = self._load_kb_metadata(kb_id)
            name = metadata.get("name", kb_id) if metadata else kb_id
            description = metadata.get("description", "") if metadata else ""

            kb = KnowledgeBase(
                id=kb_id,
                name=name,
                description=description,
                chunk_count=info.points_count,
            )
            kbs.append(kb)

        return kbs

    async def delete_knowledge_base(self, kb_id: str) -> bool:
        """Delete a knowledge base and all its data"""
        collection_name = self._get_collection_name(kb_id)
        try:
            self.client.delete_collection(collection_name)
            self._kb_cache.pop(kb_id, None)
            # Also delete metadata
            try:
                import uuid as uuid_module
                point_uuid = str(uuid_module.uuid5(uuid_module.NAMESPACE_DNS, kb_id))
                self.client.delete(
                    collection_name=self.METADATA_COLLECTION,
                    points_selector=qmodels.PointIdsList(
                        points=[point_uuid],
                    ),
                )
            except Exception:
                pass
            return True
        except Exception:
            return False

    async def update_knowledge_base(
        self,
        kb_id: str,
        name: str = None,
        description: str = None,
        metadata: Dict[str, Any] = None
    ) -> Optional[KnowledgeBase]:
        """Update knowledge base metadata"""
        kb = await self.get_knowledge_base(kb_id)
        if not kb:
            return None

        # Update fields
        if name is not None:
            kb.name = name
        if description is not None:
            kb.description = description
        if metadata is not None:
            kb.metadata.update(metadata)
        kb.updated_at = datetime.now()

        # Save updated metadata
        self._save_kb_metadata(kb)
        self._kb_cache[kb_id] = kb

        return kb

    async def get_knowledge_base(self, kb_id: str) -> Optional[KnowledgeBase]:
        """Get a single knowledge base by ID"""
        # Check cache first
        if kb_id in self._kb_cache:
            return self._kb_cache[kb_id]

        # Try to load from metadata collection
        try:
            import uuid as uuid_module
            point_uuid = str(uuid_module.uuid5(uuid_module.NAMESPACE_DNS, kb_id))

            results = self.client.retrieve(
                collection_name=self.METADATA_COLLECTION,
                ids=[point_uuid],
            )

            if results:
                payload = results[0].payload
                kb = KnowledgeBase(
                    id=kb_id,
                    name=payload.get("name", kb_id),
                    description=payload.get("description", ""),
                    created_at=datetime.fromisoformat(payload["created_at"]) if payload.get("created_at") else datetime.now(),
                    updated_at=datetime.fromisoformat(payload["updated_at"]) if payload.get("updated_at") else datetime.now(),
                    metadata=payload.get("metadata", {}),
                )
                self._kb_cache[kb_id] = kb
                return kb
        except Exception:
            pass

        return None

    # ==================== Document Management ====================

    async def add_document(
        self,
        doc: Document,
        kb_id: str = "default"
    ) -> List[str]:
        """
        Add a document to a knowledge base

        Returns list of chunk IDs
        """
        # Ensure collection exists
        self._ensure_collection(kb_id)

        chunks = self._chunk_text(doc.content)
        chunk_ids = []

        points = []
        for i, chunk in enumerate(chunks):
            chunk_id = self._generate_id(chunk, doc.source)
            chunk_ids.append(chunk_id)

            # Generate embedding
            embedding = self.encoder.encode(chunk).tolist()

            # Create point
            points.append(qmodels.PointStruct(
                id=chunk_id,
                vector=embedding,
                payload={
                    "content": chunk,
                    "title": doc.title,
                    "source": doc.source,
                    "source_type": doc.source_type,
                    "url": doc.url,
                    "file_type": doc.file_type,
                    "doc_id": doc.id,
                    "kb_id": kb_id,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "created_at": doc.created_at.isoformat(),
                    **doc.metadata,
                },
            ))

        # Upsert to Qdrant
        self.client.upsert(
            collection_name=self._get_collection_name(kb_id),
            points=points,
        )

        return chunk_ids

    async def add_file(
        self,
        file_content: bytes,
        filename: str,
        kb_id: str = "default",
        source: str = "",
        url: str = ""
    ) -> Dict[str, Any]:
        """Parse and add a file to knowledge base"""
        # Parse file
        result = FileParser.parse(file_content, filename)

        if not result.success:
            return {
                "success": False,
                "error": result.error
            }

        # Create document
        doc = Document(
            id=str(uuid.uuid4()),
            content=result.content,
            title=result.title,
            source=source or filename,
            source_type="file",
            url=url,
            file_type=Path(filename).suffix.lower(),
            file_size=len(file_content),
            kb_id=kb_id,
            metadata=result.metadata
        )

        # Add to knowledge base
        chunk_ids = await self.add_document(doc, kb_id)

        return {
            "success": True,
            "doc_id": doc.id,
            "title": doc.title,
            "chunks": len(chunk_ids),
            "file_type": doc.file_type,
        }

    async def add_webpage(
        self,
        url: str,
        kb_id: str = "default"
    ) -> Dict[str, Any]:
        """Fetch and add a webpage to knowledge base"""
        # Parse URL
        result = FileParser.parse_url(url)

        if not result.success:
            return {
                "success": False,
                "error": result.error
            }

        # Create document
        doc = Document(
            id=str(uuid.uuid4()),
            content=result.content,
            title=result.title,
            source=url,
            source_type="webpage",
            url=url,
            file_type="html",
            kb_id=kb_id,
            metadata=result.metadata
        )

        # Add to knowledge base
        chunk_ids = await self.add_document(doc, kb_id)

        return {
            "success": True,
            "doc_id": doc.id,
            "title": doc.title,
            "chunks": len(chunk_ids),
            "url": url,
        }

    # ==================== Search ====================

    async def search(
        self,
        query: str,
        kb_id: str = "default",
        limit: int = 5,
        score_threshold: float = 0.5,
    ) -> List[SearchResult]:
        """Search within a knowledge base"""
        collection_name = self._get_collection_name(kb_id)

        # Check if collection exists
        collections = self.client.get_collections().collections
        if not any(c.name == collection_name for c in collections):
            return []

        # Generate query embedding
        query_embedding = self.encoder.encode(query).tolist()

        # Search
        results = self.client.query_points(
            collection_name=collection_name,
            query=query_embedding,
            limit=limit,
            score_threshold=score_threshold,
        )

        # Convert to SearchResult
        search_results = []
        for hit in results.points:
            payload = hit.payload
            doc = Document(
                id=payload.get("doc_id", ""),
                content=payload.get("content", ""),
                title=payload.get("title", ""),
                source=payload.get("source", ""),
                source_type=payload.get("source_type", ""),
                url=payload.get("url", ""),
                file_type=payload.get("file_type", ""),
                kb_id=payload.get("kb_id", kb_id),
                metadata={k: v for k, v in payload.items() if k not in [
                    "content", "title", "source", "source_type", "url",
                    "file_type", "doc_id", "kb_id"
                ]},
            )
            search_results.append(SearchResult(
                document=doc,
                score=hit.score,
                chunk_index=payload.get("chunk_index", 0)
            ))

        return search_results

    async def search_all(
        self,
        query: str,
        limit: int = 5,
        score_threshold: float = 0.5,
    ) -> Dict[str, List[SearchResult]]:
        """Search across all knowledge bases"""
        kbs = await self.list_knowledge_bases()
        results = {}

        for kb in kbs:
            kb_results = await self.search(query, kb.id, limit, score_threshold)
            if kb_results:
                results[kb.id] = kb_results

        return results

    async def build_context(
        self,
        query: str,
        kb_id: str = "default",
        max_tokens: int = 2000,
        limit: int = 5,
    ) -> tuple[str, List[Dict[str, str]]]:
        """Build context for LLM from relevant documents"""
        results = await self.search(query, kb_id, limit=limit)

        if not results:
            return "", []

        # Filter out low-relevance results
        results = [r for r in results if r.score >= 0.5]

        if not results:
            return "", []

        context_parts = []
        sources = []
        total_length = 0
        seen_sources = set()

        for result in results:
            doc = result.document
            content = doc.content

            # Estimate tokens (rough: 1 token ≈ 2 chars for Chinese)
            content_tokens = len(content) // 2

            if total_length + content_tokens > max_tokens:
                remaining = (max_tokens - total_length) * 2
                if remaining > 100:
                    content = content[:remaining] + "..."
                else:
                    break

            context_parts.append(f"【来源: {doc.source}】\n{content}")

            # Deduplicate sources
            source_key = f"{doc.title}:{doc.url}"
            if source_key not in seen_sources:
                seen_sources.add(source_key)
                sources.append({
                    "title": doc.title or doc.source,
                    "url": doc.url,
                    "source": doc.source,
                    "score": round(result.score, 3),
                })

            total_length += content_tokens

        context = "\n\n---\n\n".join(context_parts)
        return context, sources

    # ==================== Document Listing & Management ====================

    async def list_documents(
        self,
        kb_id: str = "default"
    ) -> List[Dict[str, Any]]:
        """List all documents in a knowledge base"""
        collection_name = self._get_collection_name(kb_id)

        # Check if collection exists
        collections = self.client.get_collections().collections
        if not any(c.name == collection_name for c in collections):
            return []

        all_docs: Dict[str, Dict[str, Any]] = {}
        offset = None

        while True:
            result = self.client.scroll(
                collection_name=collection_name,
                limit=100,
                offset=offset,
                with_payload=True,
                with_vectors=False,
            )

            points, next_offset = result

            for point in points:
                p = point.payload
                doc_id = p.get("doc_id", "")
                if doc_id not in all_docs:
                    all_docs[doc_id] = {
                        "doc_id": doc_id,
                        "title": p.get("title", ""),
                        "source": p.get("source", ""),
                        "source_type": p.get("source_type", ""),
                        "url": p.get("url", ""),
                        "file_type": p.get("file_type", ""),
                        "chunk_count": 0,
                        "preview": "",
                        "created_at": p.get("created_at", ""),
                    }
                all_docs[doc_id]["chunk_count"] += 1
                if not all_docs[doc_id]["preview"] and p.get("content"):
                    all_docs[doc_id]["preview"] = p["content"][:150]

            if next_offset is None:
                break
            offset = next_offset

        return list(all_docs.values())

    async def delete_document(self, kb_id: str, doc_id: str) -> bool:
        """Delete a document from a knowledge base"""
        collection_name = self._get_collection_name(kb_id)

        try:
            self.client.delete(
                collection_name=collection_name,
                points_selector=qmodels.FilterSelector(
                    filter=qmodels.Filter(
                        must=[
                            qmodels.FieldCondition(
                                key="doc_id",
                                match=qmodels.MatchValue(value=doc_id),
                            )
                        ]
                    )
                ),
            )
            return True
        except Exception:
            return False

    async def get_stats(self, kb_id: str = "default") -> Dict[str, Any]:
        """Get knowledge base statistics"""
        collection_name = self._get_collection_name(kb_id)

        # Check if collection exists
        collections = self.client.get_collections().collections
        if not any(c.name == collection_name for c in collections):
            return {"total_documents": 0, "status": "not_found"}

        info = self.client.get_collection(collection_name)

        # Count unique documents
        docs = await self.list_documents(kb_id)

        return {
            "total_chunks": info.points_count,
            "total_documents": len(docs),
            "status": str(info.status),
        }


# Singleton instance
_kb_manager: Optional[KnowledgeBaseManager] = None


def get_kb_manager() -> KnowledgeBaseManager:
    """Get or create knowledge base manager instance"""
    global _kb_manager
    if _kb_manager is None:
        _kb_manager = KnowledgeBaseManager()
    return _kb_manager
